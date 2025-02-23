import os
from docx import Document
from win32com import client
import pandas as pd

# 讀取 Excel 圖表
df = pd.read_excel("圖表.xlsx")

# 建立 Word 報告
doc = Document()
doc.add_heading("自動化報告", level=1)
doc.add_paragraph("以下是來自 Excel 的數據摘要:")

table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
for j, col in enumerate(df.columns):
    table.cell(0, j).text = col
for i, row in df.iterrows():
    for j, value in enumerate(row):
        table.cell(i + 1, j).text = str(value)

doc.save("報告.docx")

# 轉換成 PDF
word = client.Dispatch("Word.Application")
doc = word.Documents.Open(os.path.abspath("報告.docx"))
doc.SaveAs(os.path.abspath("報告.pdf"), FileFormat=17)
doc.Close()
word.Quit()

print("已生成報告並輸出為 PDF")
