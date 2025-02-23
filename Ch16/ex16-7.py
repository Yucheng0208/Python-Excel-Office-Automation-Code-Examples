import pandas as pd
from docx import Document

df = pd.read_excel("成績管理 4.xlsx")
doc = Document("個人履歷模板.docx")

doc.add_heading("學生成績", level=2)
table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

# 添加標題行
for j, col in enumerate(df.columns):
    table.cell(0, j).text = col

# 填入數據
for i, row in df.iterrows():
    for j, value in enumerate(row):
        table.cell(i + 1, j).text = str(value)

doc.save("成績報告.docx")
print("已建立成績報告 Word 文件")
