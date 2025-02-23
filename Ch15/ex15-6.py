from docx import Document

doc = Document()
doc.add_heading("個人履歷", level=1)

doc.add_paragraph("姓名: 王小明")
doc.add_paragraph("電話: 0987-654-321")
doc.add_paragraph("Email: example@email.com")
doc.add_paragraph("技能: Python, AI, 數據分析")

doc.save("個人履歷.docx")
print("已建立個人履歷.docx")
